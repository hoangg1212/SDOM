import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
from docx import Document

class App:
    def __init__(self, master):
        self.master = master
        self.master.title("Thông tin đăng ký")
        
        # Tạo các form
        self.create_personal_info_form()
        self.create_registration_status_form()
        self.create_terms_form()

        # Nút Enter và Thoát
        self.button_frame = tk.Frame(master)
        self.button_frame.pack(pady=10)
        
        self.submit_button = tk.Button(self.button_frame, text="Enter", command=self.show_info)
        self.submit_button.pack(side=tk.LEFT, padx=5)
        
        self.exit_button = tk.Button(self.button_frame, text="Thoát", command=master.quit)
        self.exit_button.pack(side=tk.LEFT, padx=5)

        # Tạo tài liệu Word
        self.doc = Document()

    def create_personal_info_form(self):
        self.personal_frame = tk.LabelFrame(self.master, text="Thông tin cá nhân", padx=10, pady=10)
        self.personal_frame.pack(padx=10, pady=10, fill="both", expand="yes")

        # Họ, Tên, Giới tính
        tk.Label(self.personal_frame, text="Họ:").grid(row=0, column=0)
        self.last_name_entry = tk.Entry(self.personal_frame)
        self.last_name_entry.grid(row=0, column=1)

        tk.Label(self.personal_frame, text="Tên:").grid(row=0, column=2)
        self.first_name_entry = tk.Entry(self.personal_frame)
        self.first_name_entry.grid(row=0, column=3)

        tk.Label(self.personal_frame, text="Giới tính:").grid(row=0, column=4)
        self.gender_var = tk.StringVar(value="Mr")
        ttk.Combobox(self.personal_frame, textvariable=self.gender_var, values=["Mr.", "Mrs.", "Miss", "Dr.", "Prof."]).grid(row=0, column=5)

        # Tuổi và Quốc tịch
        tk.Label(self.personal_frame, text="Tuổi:").grid(row=1, column=0)
        self.age_spinbox = tk.Spinbox(self.personal_frame, from_=0, to=120, width=5)
        self.age_spinbox.grid(row=1, column=1)

        tk.Label(self.personal_frame, text="Quốc tịch:").grid(row=1, column=2)
        self.nationality_entry = tk.Entry(self.personal_frame)
        self.nationality_entry.grid(row=1, column=3)

    def create_registration_status_form(self):
        self.registration_frame = tk.LabelFrame(self.master, text="Tình trạng đăng ký", padx=10, pady=10)
        self.registration_frame.pack(padx=10, pady=10, fill="both", expand="yes")

        # Tick Đã đăng ký và Số lượng khóa học
        self.registered_var = tk.BooleanVar()
        tk.Checkbutton(self.registration_frame, text="Đã đăng ký", variable=self.registered_var).grid(row=0, column=0)

        tk.Label(self.registration_frame, text="Số lượng khóa học hoàn thành:").grid(row=0, column=1)
        self.completed_courses_spinbox = tk.Spinbox(self.registration_frame, from_=0, to=100, width=5)
        self.completed_courses_spinbox.grid(row=0, column=2)

        # Số lượng học kỳ
        tk.Label(self.registration_frame, text="Số lượng học kỳ:").grid(row=1, column=1)
        self.semesters_spinbox = tk.Spinbox(self.registration_frame, from_=0, to=10, width=5)
        self.semesters_spinbox.grid(row=1, column=2)

    def create_terms_form(self):
        self.terms_frame = tk.LabelFrame(self.master, text="Điều khoản và điều kiện", padx=10, pady=10)
        self.terms_frame.pack(padx=10, pady=10, fill="both", expand="yes")

        self.terms_var = tk.BooleanVar()
        tk.Checkbutton(self.terms_frame, text="Đồng ý chấp nhận điều khoản", variable=self.terms_var).grid(row=0, columnspan=2)

    def show_info(self):
        # Kiểm tra điều kiện
        if not self.terms_var.get():
            messagebox.showwarning("Cảnh báo", "Vui lòng đồng ý điều khoản và điều kiện.")
            return
        
        info = f"Họ: {self.last_name_entry.get()}\n"
        info += f"Tên: {self.first_name_entry.get()}\n"
        info += f"Giới tính: {self.gender_var.get()}\n"
        info += f"Tuổi: {self.age_spinbox.get()}\n"
        info += f"Quốc tịch: {self.nationality_entry.get()}\n"
        info += f"Đã đăng ký: {'Có' if self.registered_var.get() else 'Không'}\n"
        info += f"Số lượng khóa học hoàn thành: {self.completed_courses_spinbox.get()}\n"
        info += f"Số lượng học kỳ: {self.semesters_spinbox.get()}\n"

        messagebox.showinfo("Thông tin đã nhập", info)
        
        # Lưu vào file Word
        self.save_to_word(info)

    def save_to_word(self, info):
        try:
            self.doc.add_heading('Thông tin đăng ký', level=1)
            self.doc.add_paragraph(info)
            self.doc.save('thong_tin_dang_ky.docx')
            messagebox.showinfo("Thông báo", "Thông tin đã được lưu vào file 'thong_tin_dang_ky.docx'")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể lưu thông tin: {e}")

# Khởi tạo ứng dụng
root = tk.Tk()
app = App(root)
root.mainloop()
